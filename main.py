import qrcode
from PIL import Image
import cv2 as cv
from wand.image import Image as wandImage
from wand.display import display
from collections import defaultdict
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PyPDF2 import PdfFileReader, PdfFileWriter
import os, docx, json
from docx.shared import Cm
import shutil
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pympler.tracker import SummaryTracker

def fullname_convert(name):
    return name.replace(' ', '_')

class QRCodeMaker:
    
    def __init__(self, errorCorrect=qrcode.ERROR_CORRECT_H, cls=None, size=(120, 120)):
        self._QRCode = qrcode.QRCode(error_correction=errorCorrect)
        self.size = size
        self._QRCodeImage = None
        self.tempFilenameTemplate = "temp/QRcode{}_{}.png"

    def make(self, page):
        self._QRCode.add_data("{}:{}:{}".format(self.ID, self.cls, page))
        self._QRCodeImage = self._QRCode.make_image()
        self._QRCode.clear()
        self._QRCodeImage = self._QRCodeImage.resize(self.size)
        filename = self.tempFilenameTemplate.format(self.ID, page)
        print(filename)
        self._QRCodeImage.save(filename)
        return filename

class WordHeaderChanger:

    def __init__(self, desc_filename, output_filename="temp/output{}.docx"):
        self.docs_count = 0
        self.output_filename = output_filename
        self.QRMaker = QRCodeMaker()
        with open(desc_filename, 'r') as desc_file:
            self.docsByCls = {cls: docx.Document(filename) for cls, filename in json.load(desc_file).items()}

    def update_doc(self, doc, ID):
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text[:2] == "ID":
                    paragraph.text = "ID: {}".format(ID)
        doc.add_picture(self.QRMaker.tempFilenameTemplate, width=Cm(3))
    
    def add_doc(self, doc):
        doc.save(self.output_filename.format(self.docs_count))
        self.docs_count += 1

    def add_work(self, cls):
        self.QRMaker.cls = cls
        curID = self.QRMaker.ID
        QRImage = self.QRMaker.make()
        self.update_doc(self.docsByCls[cls], curID)
        self.add_doc(self.docsByCls[cls])


class PdfHeaderChanger:
    def __init__(self, desc_filename, output_path="output", max_count=100, startingDocsId=1, startingStudentId=1, double_sided_print=True):
        self.docs_count = startingDocsId - 1
        self.page_count = 0
        self.workID = startingStudentId
        self.max_count = max_count
        self.output_filename = '{}/to_print'.format(output_path) + '{}.pdf'
        self.outputPdf = None
        self.input_dir_prefix = "input//"
        self.QRMaker = QRCodeMaker()
        self.ds_print = double_sided_print
        self.blankFileName = self.input_dir_prefix + "blank.pdf"
        self.additionalFileName = self.input_dir_prefix + "additional24.pdf"
        if desc_filename:
            with open(desc_filename, 'r') as desc_file:
                self.docsByCls = json.load(desc_file)
        pdfmetrics.registerFont(TTFont("DejaVuSans", "DejaVuSans.ttf"))

        self.newOutputDocument()
        
        if not os.path.isdir('temp'):
            os.mkdir('temp')

    def newOutputDocument(self):
        if self.outputPdf:
            if self.outputPdf.getNumPages() > 0:
                with open(self.output_filename.format(self.docs_count), "wb") as out:
                    self.outputPdf.write(out)
            for file in self.toClose:
                file.close()
        self.toClose = []
        self.outputPdf = PdfFileWriter()
        self.docs_count += 1
        self.page_count = 0

    def make_pdf2merge(self, page_number_for_student, custom_page_count=None):
        page_count = self.page_count
        if custom_page_count is not None:
            page_count = custom_page_count
        QRfilename = self.QRMaker.make(page_count)
        filename = 'temp/pdf2merge{}_{}.pdf'.format(self.docs_count, page_count)
        newPdf = canvas.Canvas(filename)
        newPdf.setFont("Times-Roman", 12)
        newPdf.drawImage(QRfilename, 462, 712)
        newPdf.drawString(50, 806,"ID: {}".format(self.id_to_print()))
        newPdf.setFont("DejaVuSans", 12)
        newPdf.drawString(50, 26,"стр. {}".format(page_number_for_student))
        newPdf.save()
        return filename
    
    def id_to_print(self):
        return "25" + self.cls_id + str(self.workID)

    def make_title2merge(self, cls):
        filename = 'temp/title2merge{}_{}.pdf'.format(self.docs_count, self.page_count)
        newPdf = canvas.Canvas(filename)
        newPdf.setFont("Times-Roman", 12)
        newPdf.drawString(236, 673,"{}".format(self.id_to_print()))
        newPdf.drawString(381, 604,"{}".format(cls))
        newPdf.drawString(236, 233,"{}".format(self.id_to_print()))
        newPdf.drawString(381, 167,"{}".format(cls))
        newPdf.save()
        return filename

    def openPdfReader(self, filename):
        stream = open(filename, 'rb')
        self.toClose.append(stream)
        return PdfFileReader(stream)

    def mergeAddPages(self, pdfs, pages):
        curPage = pdfs[0].getPage(pages[0])
        for i in range(1, len(pdfs)):
            curPage.mergePage(pdfs[i].getPage(pages[i]))
        self.outputPdf.addPage(curPage)
        
    def add_blank_page(self):
        blank_page = self.openPdfReader(self.blankFileName)
        self.outputPdf.addPage(blank_page.getPage(0))

    def add_additional_page(self, qr_code_pdf):
        self.mergeAddPages((self.additional_page, qr_code_pdf), (0, 0))


    def add_empty_QR_Id(self, cls, page_count, custom_page_number=200):
        self.QRMaker.cls = cls
        self.cls_id = str(cls)
        self.QRMaker.ID = self.workID

        for page_num in range(page_count):
            page_filename = self.make_pdf2merge(custom_page_number)
            qrcodePdf = self.openPdfReader(page_filename)
            additional_page = self.openPdfReader(self.additionalFileName)
            self.mergeAddPages((additional_page, qrcodePdf), (0, 0))
            self.page_count += 1
            custom_page_number += 1
        if self.ds_print and page_count % 2 == 1:
            assert "Blank page detected"
            self.add_blank_page()

        self.workID += 1
        if self.page_count > self.max_count:
            self.newOutputDocument()
        

    def add_work(self, cls, cls_key):
        self.QRMaker.cls = cls
        self.cls_id = str(cls)
        if cls != cls_key:
            self.cls_id += "3"
        self.QRMaker.ID = self.workID

        title_filename = self.make_title2merge(cls)
        titleText = self.openPdfReader(title_filename)
        titlePdf = self.openPdfReader(self.input_dir_prefix + self.docsByCls['title'])

        #adding title to every work
        self.mergeAddPages((titlePdf, titleText), (0, 0))
        self.outputPdf.addBlankPage()



        current_numeration = 1
        for cls_filename in sorted(self.docsByCls[cls_key], key=lambda s: 'test' in s, reverse=True):
            
            taskPdf = self.openPdfReader(self.input_dir_prefix + cls_filename)

            page_count = taskPdf.getNumPages()
            for page_num in range(page_count):
                #make QRCode
                pdf2merge_filename = self.make_pdf2merge(current_numeration)
                current_numeration += 1
                qrcodePdf = self.openPdfReader(pdf2merge_filename)
                # merge current page
                self.mergeAddPages((taskPdf, qrcodePdf), (page_num, 0))
                self.page_count += 1
            if self.ds_print and page_count % 2 == 1:
                self.add_blank_page()

        self.workID += 1
        if self.page_count > self.max_count:
            self.newOutputDocument()
    
    def finish(self):
        self.newOutputDocument()
        
        if os.path.isdir('temp'):
            try:
                shutil.rmtree('temp')
            except OSError as e:
                print("Error: %s - %s." % (e.filename, e.strerror))

def getPageNumberFromFile(filename):
    prefix = filename.split('.')[0]
    page_number = prefix.split('_')[-1]
    return int(page_number)


class QRCodeScanner:

    def __init__(self, output_folder, load=None, save=None):
        self._detector = cv.QRCodeDetector()
        if save:
            self.save_filename = "save"
        if load:
            self.loadLogs("save")
        else:
            self._visited = set()
        self._image_type = 'png'
        self._output_path = output_folder
        self._inter_output_path = "onelistpdfs"
        self.fileCount = defaultdict(int, {"5": 16, "7": 32, "8": 25, "9": 42})
        
        self.memoryTracker = SummaryTracker()
        
        if not os.path.isdir(self._inter_output_path):
            os.mkdir(self._inter_output_path)

        if not os.path.isdir('blacklist'):
            os.mkdir('blacklist')
        
        if not os.path.isdir('temp'):
            os.mkdir('temp')

    def __del__(self):
        if os.path.isdir('temp'):
            try:
                shutil.rmtree('temp')
            except OSError as e:
                print("Error: %s - %s." % (e.filename, e.strerror))

    
    def cropNoiseRemove(self, wImage):
        w, h = wImage.size
        wImage.crop(int(w * 0.77), int(h * 0.0), int(w ), int(h * 0.21))
        #wImage.statistic('mean', width=5, height=5)
        wImage.brightness_contrast(int(0), int(50))

    def rotateCropNoiseRemove(self, wImage):
        wImage.rotate(180)
        w, h = wImage.size
        wImage.crop(int(w * 0.77), int(h * 0.0), int(w ), int(h * 0.21))
        #wImage.statistic('mean', width=5, height=5)
        wImage.brightness_contrast(int(0), int(50))

    def extractQrCode(self, image, image_filename, preprocessFunc):
        with wandImage(image) as wImage:
            preprocessFunc(wImage)
            wImage.save(filename=image_filename)
            qr_string, points, straight_qrcode = self._detector.detectAndDecode(cv.imread(image_filename))
        return qr_string

    def printToFile(self, out_file, page, rotated):
        pdfWriter = PdfFileWriter()
        with open(out_file, 'wb') as out:
            if rotated:
                page.rotateClockwise(180)
            pdfWriter.addPage(page)
            pdfWriter.write(out)

    def makeFilename(self, id, cls, page):
        idpath = os.path.join(self._inter_output_path, id)
        if not os.path.isdir(idpath):
            os.mkdir(idpath)
        return os.path.join(idpath, "{}_{}_{}.pdf".format(id, cls, page))


    def tryParse(self, source, pdf_filename, pdfReader, try_rotate=False, bad_processing=False):
        if pdf_filename in self._visited:
            return
        for i, image in enumerate(source.sequence):
            if bad_processing:
                with wandImage(image) as img:
                    img.save(filename='temp/show_image.jpg')
                    while True:
                        try:
                            display(img)
                            input_string = input("type 'bad' for empty doc, else 'id:cls:page:rotated'").strip()
                            if input_string != 'bad':
                                id, cls, page, rotated = input_string.split(":")
                                rotated = bool(rotated)
                                out_file = self.makeFilename(id, cls, page)
                                self.printToFile(out_file, pdfReader.getPage(i), rotated)
                            break
                        except ValueError:
                            continue
                continue
            image_filename = 'temp/scan{}_{}.jpeg'.format(i, hash(os.path.basename(pdf_filename)))
            rotated_image_filaname = 'temp/scan{}_{}_rotated.jpeg'.format(i, hash(os.path.basename(pdf_filename)))
            qr_string = self.extractQrCode(image, image_filename, self.cropNoiseRemove)
            rotated = False
            if not qr_string and try_rotate:
                qr_string = self.extractQrCode(image, rotated_image_filaname, self.rotateCropNoiseRemove)
                rotated = True
            if qr_string:
                print(qr_string)
                id, cls, page = qr_string.split(":")
                out_file = self.makeFilename(id, cls, page)
            if not qr_string:
                print("can't find QRcode in file {}".format(pdf_filename))
                out_file = os.path.join("blacklist", "page_{}_{}".format(i, os.path.basename(pdf_filename)))
            self.printToFile(out_file, pdfReader.getPage(i), rotated)
            
        self._visited.add(pdf_filename)
        
    def scanPdf(self, pdf_filename, resolution=300, track_memory=False, **kwargs):
        if track_memory:
            self.memoryTracker.print_diff()
        with (wandImage(filename=pdf_filename, resolution=resolution)) as source, open(pdf_filename, 'rb') as pdfStream:
            pdfReader = PdfFileReader(pdfStream)
            self.tryParse(source, pdf_filename, pdfReader, **kwargs)

    
    def collectWorks(self):
        print('collecting works')
        for path in os.listdir(self._inter_output_path):
            dir = os.path.join(self._inter_output_path, path)
            if os.path.isdir(dir):
                self.collectWorksFromPath(dir)
        if self.save_filename == "save":
            self.saveLogs()

    def collectWorksFromPath(self, dir):
        pdfWriter = PdfFileWriter()
        toClose = []
        counter = 0
        min_amount = 9999999
        max_amount = 0
        for path in os.listdir(dir):
            file = os.path.join(dir, path)
            if os.path.isfile(file):
                counter += 1
                page_number = getPageNumberFromFile(file)
                min_amount = min(min_amount, page_number)
                max_amount = max(max_amount, page_number)
                file = open(file, 'rb')
                id, cls, page = path.split('_')
                toClose.append((page, id, cls, file))
        toClose.sort()
        if (counter < self.fileCount[cls]):
            print("wrong file count for ID {}, cls {}, expected {}, got {}".format(id, cls, self.fileCount[cls], counter))
        if (max_amount - min_amount + 1 != counter):
            print("wrong file count for ID {}, cls {}, max_page_number {}, min_page_number {}, got {}".format(id, cls, max_amount, min_amount, counter))
        for page, id, cls, file in toClose:
            pdfReader = PdfFileReader(file)
            pdfWriter.addPage(pdfReader.getPage(0))
        with open(os.path.join(self._output_path, "{}cls_{}id.pdf".format(cls, id)), 'wb') as out:
            pdfWriter.write(out)
        for page, id, cls, file in toClose:
            file.close()

    def checkBlacklist(self):
        pass
        

    def saveLogs(self):
        with open(self.save_filename, "w") as fw:
            json.dump({"visited": list(self._visited)}, fw)

    def loadLogs(self, filename):
        with open(filename, "w") as fw:
            json_dict = json.load(fw)
            self._visited = set(json_dict["visited"])
    
        
if __name__ == "__main__":
    pdf_ch = PdfHeaderChanger("cls_files.json")
    for i in range(900):
        pdf_ch.add_work("5")
    pdf_ch.finish()
    '''scanner = QRCodeScanner()
    scanner.scanPdf("output/1.pdf")
    scanner.scanPdf("output/2.pdf")
    scanner.collectWorks()
    scanner.printWorks()'''